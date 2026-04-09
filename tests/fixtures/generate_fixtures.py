"""
generate_fixtures.py
Generates small test files for clio-tools integration tests.
Run once before running tests: python tests/fixtures/generate_fixtures.py

Outputs:
    tests/fixtures/sample.pdf      – 2-page scanned-style PDF
    tests/fixtures/sample.mp3      – 5-second speech clip
    tests/fixtures/sample.jpg      – simple test image
    tests/fixtures/sample.docx     – short Word document
    tests/fixtures/sample.txt      – plain text file
    tests/fixtures/sample.md       – markdown file
"""

import sys
import subprocess
from pathlib import Path

FIXTURES = Path(__file__).parent


def generate_pdf():
    """Creates a simple PDF with text using reportlab, or falls back to fpdf2."""
    out = FIXTURES / "sample.pdf"
    if out.exists():
        print(f"  EXISTS: {out.name}")
        return

    try:
        from reportlab.pdfgen import canvas
        c = canvas.Canvas(str(out))
        c.setFont("Helvetica", 14)
        c.drawString(72, 750, "Clio Tools – Test Document")
        c.drawString(72, 720, "This is page 1 of the test fixture.")
        c.drawString(72, 690, "Used for clio-docs OCR testing.")
        c.showPage()
        c.drawString(72, 750, "Page 2")
        c.drawString(72, 720, "Second page content for multi-page testing.")
        c.save()
        print(f"  CREATED: {out.name} (reportlab)")
    except ImportError:
        try:
            from fpdf import FPDF
            pdf = FPDF()
            pdf.add_page()
            pdf.set_font("Helvetica", size=14)
            pdf.cell(0, 10, "Clio Tools - Test Document", ln=True)
            pdf.cell(0, 10, "This is page 1 of the test fixture.", ln=True)
            pdf.add_page()
            pdf.cell(0, 10, "Page 2", ln=True)
            pdf.output(str(out))
            print(f"  CREATED: {out.name} (fpdf2)")
        except ImportError:
            print(f"  SKIP: {out.name} – install reportlab or fpdf2")


def generate_mp3():
    """Creates a short MP3 using edge-tts."""
    out = FIXTURES / "sample.mp3"
    if out.exists():
        print(f"  EXISTS: {out.name}")
        return
    try:
        import asyncio, edge_tts
        text  = "This is a test audio file for clio-transcribe. Clio tools test."
        voice = "en-GB-SoniaNeural"

        async def _gen():
            communicate = edge_tts.Communicate(text, voice)
            await communicate.save(str(out))

        asyncio.run(_gen())
        print(f"  CREATED: {out.name} (edge-tts)")
    except ImportError:
        print(f"  SKIP: {out.name} – install edge-tts")
    except Exception as e:
        print(f"  FAIL: {out.name} – {e}")


def generate_jpg():
    """Creates a simple test image using Pillow."""
    out = FIXTURES / "sample.jpg"
    if out.exists():
        print(f"  EXISTS: {out.name}")
        return
    try:
        from PIL import Image, ImageDraw, ImageFont
        img  = Image.new("RGB", (400, 300), color=(240, 240, 240))
        draw = ImageDraw.Draw(img)
        draw.rectangle([20, 20, 380, 280], outline=(100, 100, 100), width=2)
        draw.text((40, 60),  "Clio Tools",        fill=(50, 50, 50))
        draw.text((40, 100), "Test Image",         fill=(50, 50, 50))
        draw.text((40, 140), "clio-vision fixture",fill=(80, 80, 80))
        draw.text((40, 180), "2026-03-29",         fill=(80, 80, 80))
        img.save(str(out), "JPEG", quality=85)
        print(f"  CREATED: {out.name} (Pillow)")
    except ImportError:
        print(f"  SKIP: {out.name} – install Pillow")


def generate_docx():
    """Creates a short Word document using python-docx."""
    out = FIXTURES / "sample.docx"
    if out.exists():
        print(f"  EXISTS: {out.name}")
        return
    try:
        from docx import Document
        doc = Document()
        doc.add_heading("Clio Tools Test Document", 0)
        doc.add_paragraph("This is a test document for clio-narrate.")
        doc.add_paragraph("It contains two short paragraphs of English text.")
        doc.add_paragraph("The quick brown fox jumps over the lazy dog.")
        doc.save(str(out))
        print(f"  CREATED: {out.name} (python-docx)")
    except ImportError:
        print(f"  SKIP: {out.name} – install python-docx")


def generate_txt():
    out = FIXTURES / "sample.txt"
    if out.exists():
        print(f"  EXISTS: {out.name}")
        return
    out.write_text(
        "Clio Tools Test File\n\n"
        "This is a plain text file used for testing clio-narrate.\n"
        "It contains simple English text across multiple lines.\n\n"
        "The quick brown fox jumps over the lazy dog.\n"
        "Pack my box with five dozen liquor jugs.\n",
        encoding="utf-8"
    )
    print(f"  CREATED: {out.name}")


def generate_md():
    out = FIXTURES / "sample.md"
    if out.exists():
        print(f"  EXISTS: {out.name}")
        return
    out.write_text(
        "# Clio Tools Test\n\n"
        "- **Source:** sample.md\n"
        "- **OCR date:** 2026-03-29\n"
        "- **Pages:** 2\n\n"
        "---\n\n"
        "## Page 1\n\n"
        "<!-- source: sample.pdf | page: 1 | ocr-date: 2026-03-29 -->\n\n"
        "This is the first page of the test markdown file.\n\n"
        "---\n\n"
        "## Page 2\n\n"
        "<!-- source: sample.pdf | page: 2 | ocr-date: 2026-03-29 -->\n\n"
        "This is the second page.\n\n"
        "---\n",
        encoding="utf-8"
    )
    print(f"  CREATED: {out.name}")


if __name__ == "__main__":
    print("Generating test fixtures...\n")
    generate_pdf()
    generate_mp3()
    generate_jpg()
    generate_docx()
    generate_txt()
    generate_md()
    print("\nDone. Run tests with: python tests/run_tests.py")
